##
import numpy as np
from github import Github
import docx
from docx import Document
import re
import urllib
import os
import PIL
import glob
import pandas as pd
from PIL import Image
from docx import Document
from docxcompose.composer import Composer

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    # r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink
##
g = Github("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")
##
for repo in g.get_user().get_repos():
    print(repo.name)

##

AA=g.get_user().get_repos()
repo = g.get_repo("Nipro-MIC/Surdial-X-Pro-Observations")
BB=repo.get_issues()

# CC=BB[10].html_url
# DD=BB[10].body

##
# Report = docx.Document('Template.docx')
# # The adjustment for header
# section = Report.sections[0]
# header = section.header
# paragraph = header.paragraphs[0]
# RepoUrl=repo.html_url
# paragraph.text = "Nipro-MIC\t"+RepoUrl
# paragraph.style = Report.styles["Header"]
#
# # The adjustment for footer
# footer = section.footer
# paragraph = footer.paragraphs[0]
# # paragraph.text = 'Nipro Europe Group Companies N.V. .\r\n Blokhuisstraat 42 - 2800 Mechelen - Belgium\r\n T: +32 15 263 664'
#
#
# paragraph.text = "\t \tBlokhuisstraat 42 - 2800 Mechelen - Belgium"
# paragraph.style = Report.styles["footer"]

#BB.totalCount
ReportNames=[]
for i in range(0,BB.totalCount):
    Report = docx.Document('Template.docx')
    # The adjustment for header
    section = Report.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    RepoUrl = repo.html_url
    paragraph.text = "Nipro-MIC\t" + RepoUrl
    paragraph.style = Report.styles["Header"]

    # The adjustment for footer
    footer = section.footer
    paragraph = footer.paragraphs[0]
    # paragraph.text = 'Nipro Europe Group Companies N.V. .\r\n Blokhuisstraat 42 - 2800 Mechelen - Belgium\r\n T: +32 15 263 664'

    paragraph.text = "\t \tBlokhuisstraat 42 - 2800 Mechelen - Belgium"
    paragraph.style = Report.styles["footer"]

    Report.add_heading(BB[i].title, 2)

    #Editing the comments
    Report.add_paragraph(BB[i].body[:BB[i].body.find('!')])

    labels='Labels:'+str(str(BB[i].labels[0].name))
    Report.add_paragraph(labels)
    Report.add_paragraph()

    State=str(BB[i].state)
    Report.add_paragraph('Issue state: '+BB[i].state)


    Report.add_heading('Photos can be found using the links below', 3)
    Links = re.findall(r'(https?://[^\s]+)', BB[i].body)
    for j in range(0,len(Links)):
         p =Report.add_paragraph('')
         add_hyperlink(p, Links[j][:-1], Links[j][:-1])
    #     # pic = Report.add_paragraph()
    #     # r = pic.add_run()
         urllib.request.urlretrieve(Links[j][:-1], "Pic.bmp")
         image = Image.open('Pic.bmp')
         [x, y] = image.size
    #     # # image.show()
    #     #  fixed_height = 100
    #     # image = Image.open('Pic.bmp')
         if x >700 or y>700:
             resized_im = image.resize((round(image.size[0] * 0.15), round(image.size[1] * 0.15)))
         else:
             resized_im = image.resize((round(image.size[0] * 1), round(image.size[1] * 1)))

         resized_im.save('resized_im.bmp')
         Report.add_picture('resized_im.bmp')

    Report.add_heading('Link to the issue on Github', 3)
    q=Report.add_paragraph('')
    add_hyperlink(q, BB[i].html_url, BB[i].html_url)

         # Report.add_heading('Revision remarks', 3)
         # Report.add_page_break()
    Report.save("ReportIssue"+str(i)+".docx")
    ReportNames.append("ReportIssue"+str(i)+".docx")

##

composed='FinalReport.docx'
result = Document(ReportNames[0])
result.add_page_break()
composer = Composer(result)

for i in range(1, len(ReportNames)):
    doc = Document(ReportNames[i])

    if i != len(ReportNames) - 1:
        doc.add_page_break()

    composer.append(doc)

composer.save(composed)

##
# string values in the list
ExcelContent=['Issue Title' ,'Issue ID','Assignees','Created at','Last updated','Status','labels','link',]

Titles = []
ID=[]
Assignees=[]
LastUpdated=[]
CreatedAt=[]
Status=[]
Labels=[]
Link=[]

for i in range(0,BB.totalCount):
    Titles.append(BB[i].title)
    ID.append(BB[i].id)
    Assignees.append(BB[i].assignees[0].login)
    CreatedAt.append(str(BB[i].created_at.date()))
    LastUpdated.append(str(BB[i].updated_at.date()))
    Status.append(BB[i].state)
    Labels.append(BB[i].labels)
    Link.append(BB[i].html_url)

for i in range (0,len(Labels)):
    for j in range (0,len(Labels[i])):
        Labels[i][j]=Labels[i][j].name
# Calling DataFrame constructor on list
# df = pd.DataFrame(data, index =Titles)
# print(df)

list_tuples = list(zip(Titles, ID,Assignees,CreatedAt,LastUpdated,Status,Labels,Link))
dframe = pd.DataFrame(list_tuples, columns=ExcelContent)
dframe.to_excel("issues.xlsx",index=False)
# dframe.to_html()
##
#Saving the report
    # Report.save("Report"+str(i)+".docx")

# print('Report is ready')

##
# urllib.request.urlretrieve('https://user-images.githubusercontent.com/101181316/160349911-a99f8a95-dae3-48bc-b9b8-4f9b21a306e6.jpg', "00000002.jpg")
# urllib.request.urlretrieve('https://user-images.githubusercontent.com/101181316/160349911-a99f8a95-dae3-48bc-b9b8-4f9b21a306e6.jpg', "00000002.jpg")
# urllib.request.urlretrieve(Links[0][:-1], "00000003.jpg")
##
# if y>x:
#             height_percent = (fixed_height / float(image.size[0]))
#             width_size = int((float(image.size[1]) * float(height_percent)))
        # if y<x:
        #     height_percent = (fixed_height / float(image.size[1]))
        #     width_size = int((float(image.size[0]) * float(height_percent)))
        #
            # image2 = image.resize((width_size, fixed_height), PIL.Image.NEAREST)
        # # image2.show()



        # r.add_picture('Pic.bmp',width=Inches(2), height=Inches(2))
##
# import requests # request img from web
# import shutil # save img locally
#
# url = Links[0]
# file_name = input('Save image as (string):') #prompt user for file_name
#
# res = requests.get(url, stream = True)
#
# if res.status_code == 200:
#     with open(file_name,'wb') as f:
#         shutil.copyfileobj(res.raw, f)
#     print('Image sucessfully Downloaded: ',file_name)
# else:
#     print('Image Couldn\'t be retrieved')
#
# ##
# from download import download
# path = download(BB[0].html_url, 'Test.jpg')
# ##
# import wget
#
# url = Links[0]
#
# wget.download(url, 'pythonLogo.png')
# ##
# import urllib2,cookielib
#
# site= Links[0]
# hdr = {'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.11 (KHTML, like Gecko) Chrome/23.0.1271.64 Safari/537.11',
#        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
#        'Accept-Charset': 'ISO-8859-1,utf-8;q=0.7,*;q=0.3',
#        'Accept-Encoding': 'none',
#        'Accept-Language': 'en-US,en;q=0.8',
#        'Connection': 'keep-alive'}
#
# req = urllib2.Request(site, headers=hdr)
# ##
# try:
#     page = urllib2.urlopen(req)
# except urllib2.HTTPError, e:
#     print e.fp.read()
#
# content = page.read()
# print(content)
# ##
# from urllib.request import Request, urlopen
# req = Request('https://user-images.githubusercontent.com/101181316/160349911-a99f8a95-dae3-48bc-b9b8-4f9b21a306e6.jpg')
# webpage = urlopen(req).read()
# W=urlopen(req)
# ##
# import shutil
# import tempfile
# import urllib.request
#
# with urllib.request.urlopen('https://user-images.githubusercontent.com/101181316/160349911-a99f8a95-dae3-48bc-b9b8-4f9b21a306e6.jpg') as response:
#     with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#         shutil.copyfileobj(response, tmp_file)
#
# with open(tmp_file.name) as html:
#     pass
# ##
# from urllib.request import urlopen
# from bs4 import BeautifulSoup
# html = urlopen('https://user-images.githubusercontent.com/101181316/160349911-a99f8a95-dae3-48bc-b9b8-4f9b21a306e6.jpg')
# ##
# bsObj = BeautifulSoup(webpage)
#
# ##
# from urllib.request import Request, urlopen
# url="https://stackoverflow.com/search?q=html+error+403"
# req = Request(url, headers={'User-Agent': 'Mozilla/5.0'})
#
# web_byte = urlopen(req).read()
#
# webpage = web_byte.decode('utf-8')
# ##
##header = section.header
# mydoc.add_header("This report is generated from Nipro-MIC Github account- Surdial-X-Pro-Observations repository ", 0)
# mydoc.add_paragraph(BB[10].body)
# mydoc.save("Report.docx")
##
# repo = g.get_repo("Surdial-X-Pro-Observations")
##
# for repo in g.get_user().get_repos():
#     open_issues = repo.get_issues(state='open')
#     repo.get_issue(number=1)
#     print(repo.name)

##
# repos = g.get_user().get_repos(474889746)

##
# g = Github(base_url="https://github.com/", login_or_token="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")
##
# for repo in g.get_user().get_repos():
#     print(repo.name)
##
# for repo in g.get_user().get_repos():
#     print(repo.name)
#     repo.edit(has_wiki=False)
#     # to see all the available attributes and methods
#     print(dir(repo))
##
# import urllib.request
#
# # image_url = j #the image on the web
# # save_name = 'my_image.jpg' #local name to be saved
# # urllib.request.urlretrieve(image_url, save_name)
# ##
# from selenium import webdriver
# import time
# ##
# # Main Function
# if __name__ == '__main__':
# ##
# from selenium import webdriver
# import time
#     # Provide the email and password
# # email = 'example@example.com'
# # password = 'password'
# #
# # options = webdriver.ChromeOptions()
# # options.add_argument("--start-maximized")
# # options.add_argument('--log-level=3')
#
# # Provide the path of chromedriver present on your system.
# driver = webdriver.Chrome(executable_path="C:\Program Files\Google\Chrome\Application\chrome.exe")
# driver.set_window_size(1920, 1080)
# driver.get(Links[0])
# time.sleep(5)
# ##
# # Send a get request to the url
# driver.get(Links[0])
# time.sleep(5)
#
# # Finds the input box by name in DOM tree to send both
# # the provided email and password in it.
# driver.find_element_by_name('user').send_keys(email)
# driver.find_element_by_name('pass').send_keys(password)
#
# # Find the signin button and click on it.
# driver.find_element_by_css_selector(
#     'button.btn.btn-green.signin-button').click()
# time.sleep(5)
#
# # Returns the list of elements
# # having the following css selector.
# container = driver.find_elements_by_css_selector(
#     'div.mdl-cell.mdl-cell--9-col.mdl-cell--12-col-phone.textBold')
#
# # Extracts the text from name,
# # institution, email_id css selector.
# name = container[0].text
# try:
#     institution = container[1].find_element_by_css_selector('a').text
# except:
#     institution = container[1].text
# email_id = container[2].text
#
# # Output Example 1
# print("Basic Info")
# print({"Name": name,
#        "Institution": institution,
#        "Email ID": email})
#
# # Clicks on Practice Tab
# driver.find_elements_by_css_selector(
#     'a.mdl-navigation__link')[1].click()
# time.sleep(5)
#
# # Selected the Container containing information
# container = driver.find_element_by_css_selector(
#     'div.mdl-cell.mdl-cell--7-col.mdl-cell--12-col-phone.\
#     whiteBgColor.mdl-shadow--2dp.userMainDiv')
#
# # Selected the tags from the container
# grids = container.find_elements_by_css_selector(
#     'div.mdl-grid')
#
# # Iterate each tag and append the text extracted from it.
# res = set()
# for grid in grids:
#     res.add(grid.text.replace('\n', ':'))
#
# # Output Example 2
# print("Practice Info")
# print(res)
#
# # Quits the driver
# # driver.close()
# # driver.quit()
# ##
# from bs4 import *
# import requests
# import os
#
#
# # CREATE FOLDER
# def folder_create(images):
#     try:
#         folder_name = input("Enter Folder Name:- ")
#         # folder creation
#         os.mkdir(folder_name)
#
#     # if folder exists with that name, ask another name
#     except:
#         print("Folder Exist with that name!")
#         folder_create()
#
#     # image downloading start
#     download_images(images, folder_name)
#
#
# # DOWNLOAD ALL IMAGES FROM THAT URL
# def download_images(images, folder_name):
#     # initial count is zero
#     count = 0
#
#     # print total images found in URL
#     print(f"Total {len(images)} Image Found!")
#
#     # checking if images is not zero
#     if len(images) != 0:
#         for i, image in enumerate(images):
#             # From image tag ,Fetch image Source URL
#
#             # 1.data-srcset
#             # 2.data-src
#             # 3.data-fallback-src
#             # 4.src
#
#             # Here we will use exception handling
#
#             # first we will search for "data-srcset" in img tag
#             try:
#                 # In image tag ,searching for "data-srcset"
#                 image_link = image["data-srcset"]
#
#             # then we will search for "data-src" in img
#             # tag and so on..
#             except:
#                 try:
#                     # In image tag ,searching for "data-src"
#                     image_link = image["data-src"]
#                 except:
#                     try:
#                         # In image tag ,searching for "data-fallback-src"
#                         image_link = image["data-fallback-src"]
#                     except:
#                         try:
#                             # In image tag ,searching for "src"
#                             image_link = image["src"]
#
#                         # if no Source URL found
#                         except:
#                             pass
#
#             # After getting Image Source URL
#             # We will try to get the content of image
#             try:
#                 r = requests.get(image_link).content
#                 try:
#
#                     # possibility of decode
#                     r = str(r, 'utf-8')
#
#                 except UnicodeDecodeError:
#
#                     # After checking above condition, Image Download start
#                     with open(f"{folder_name}/images{i + 1}.jpg", "wb+") as f:
#                         f.write(r)
#
#                     # counting number of image downloaded
#                     count += 1
#             except:
#                 pass
#
#         # There might be possible, that all
#         # images not download
#         # if all images download
#         if count == len(images):
#             print("All Images Downloaded!")
#
#         # if all images not download
#         else:
#             print(f"Total {count} Images Downloaded Out of {len(images)}")
#
#
# # MAIN FUNCTION START
# def main(url):
#     # content of URL
#     r = requests.get(url)
#
#     # Parse HTML Code
#     soup = BeautifulSoup(r.text, 'html.parser')
#
#     # find all images in URL
#     images = soup.findAll('img')
#
#     # Call folder create function
#     folder_create(images)
#
#
# # take url
# url = input('https://user-images.githubusercontent.com/101181316/160349911-a99f8a95-dae3-48bc-b9b8-4f9b21a306e6.jpg)')
#
# # CALL MAIN FUNCTION
# main(url)
# ##
# import requests
# url = 'https://user-images.githubusercontent.com/101181316/160349911-a99f8a95-dae3-48bc-b9b8-4f9b21a306e6.jpg'
# values = {'username': 'milad.ghasemi@nipro-group.com',
#           'password': 'AaBb9008654'}
#
# r = requests.post(url, data=values)
# ##
# print r.content
